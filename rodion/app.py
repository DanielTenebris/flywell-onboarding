from flask import Flask, render_template, redirect, url_for, request
from docx import *
import re

app = Flask(__name__)

text = []
files = ['templates/Московкин.docx', "templates/Ищенко.docx", 'templates/Корпоративные предложения сотрудникам Росмолодежи 10.01.23.docx']


@app.route('/', methods=['POST', 'GET'])
def index():
    data = []
    if request.method == 'POST':
        search = request.form['search']
        global text, files
        for file in files:
            for paragraph in Document(file).paragraphs:
                text.append(paragraph.text)
            for table in Document(file).tables:
                for row in table.rows:
                    for cell in row.cells:
                        if len(re.findall(search, cell.text)) > 0:
                            data.append(file.split('/')[1])
            if len(re.findall(search, '\n\n'.join(text))) > 0:
                data.append(file.split('/')[1])
            text = []
        return render_template('index.html', data=set(sorted(data)))
    else:
        return render_template('index.html')


if __name__ == '__main__':
    app.run()
